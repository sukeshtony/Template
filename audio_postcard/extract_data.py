"""
extract_data.py — Extract Audio Postcard lesson data from a Word document.

The Word document has:
  - Paragraphs: Grade headers, difficulty markers (🟢/🟡/🔴), lesson titles
  - Tables (one per lesson, 3 rows × 2 cols):
      Row 0: Overview   | <topic text>
      Row 1: Instructions | <instructions text>
      Row 2: Sample     | <sample text>

Output: postcards_data.txt — a Python-ready list of dicts for data.py
"""

from docx import Document
import re

INPUT_FILE = r"C:\Users\Sukesh\Desktop\Template\audio_postcard\Module8_Audio_Postcard.docx"
OUTPUT_FILE = "postcards_data.txt"

doc = Document(INPUT_FILE)

# ---------- Step 1: Extract lesson names, grades, difficulties from paragraphs ----------
lessons_meta = []
current_grade = None
current_difficulty = None

for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue

    # Grade detection
    grade_match = re.match(r"^Grade\s+(\d+)$", text, re.IGNORECASE)
    if grade_match:
        current_grade = int(grade_match.group(1))
        continue

    # Difficulty detection
    diff_match = re.match(r"^[🟢🟡🔴\s]*(Easy|Medium|Hard)", text, re.IGNORECASE)
    if diff_match and not text.upper().startswith("LESSON"):
        current_difficulty = diff_match.group(1).upper()
        continue

    # Lesson detection
    lesson_match = re.match(r"^Lesson\s+\d+:\s+(.+)$", text)
    if lesson_match:
        lessons_meta.append({
            "lesson_name": lesson_match.group(1).strip(),
            "grade": current_grade,
            "difficulty": current_difficulty,
        })

# ---------- Step 2: Extract topic & instructions from tables ----------
tables = doc.tables

if len(tables) != len(lessons_meta):
    print(f"⚠️  Warning: {len(lessons_meta)} lessons found but {len(tables)} tables found.")
    print("   Will pair as many as possible.")

postcards_data = []
for idx, meta in enumerate(lessons_meta):
    topic = ""
    instructions = ""

    if idx < len(tables):
        table = tables[idx]
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2:
                label = cells[0].lower()
                value = cells[1]
                if label == "overview":
                    topic = value
                elif label in ("instructions", "instruction"):
                    instructions = value

    postcards_data.append({
        "lesson_name": meta["lesson_name"],
        "grade": meta["grade"],
        "difficulty": meta["difficulty"],
        "topic": topic,
        "instructions": instructions,
    })

# ---------- Step 3: Save as Python-formatted text file ----------
with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
    f.write("postcards_data = [\n")
    for entry in postcards_data:
        # Escape any quotes in text
        name = entry["lesson_name"].replace('"', '\\"')
        topic = entry["topic"].replace('"', '\\"')
        instr = entry["instructions"].replace('"', '\\"')
        f.write("    {\n")
        f.write(f'        "lesson_name": "{name}",\n')
        f.write(f'        "grade": {entry["grade"]},\n')
        f.write(f'        "difficulty": "{entry["difficulty"]}",\n')
        f.write(f'        "topic": "{topic}",\n')
        f.write(f'        "instructions": "{instr}"\n')
        f.write("    },\n")
    f.write("]\n")

print(f"✅ Extracted {len(postcards_data)} lessons")
print(f"📄 Saved to {OUTPUT_FILE}")
